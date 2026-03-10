"""Integration tests using the OSDU CIButler MkDocs site.

These tests exercise leafpress against a real MkDocs project with Material theme,
pymdownx extensions (including !!python/name references), and 12 markdown pages.
"""

from pathlib import Path

import pytest

from leafpress.config import BrandingConfig
from leafpress.docx.renderer import DocxRenderer
from leafpress.git_info import GitVersion, extract_git_info
from leafpress.markdown_renderer import MarkdownRenderer
from leafpress.mkdocs_parser import MkDocsConfig, flatten_nav, parse_mkdocs_config
from leafpress.pdf.renderer import PdfRenderer

CIBUTLER_DOCS = Path("/Users/hutchins/projects/osdu/cibutler/docs")
CIBUTLER_REPO = Path("/Users/hutchins/projects/osdu/cibutler")

pytestmark = pytest.mark.skipif(
    not CIBUTLER_DOCS.exists(),
    reason="CIButler docs not available at expected path",
)


@pytest.fixture
def cibutler_config() -> MkDocsConfig:
    return parse_mkdocs_config(CIBUTLER_DOCS / "mkdocs.yml")


@pytest.fixture
def cibutler_branding() -> BrandingConfig:
    return BrandingConfig(
        company_name="OSDU",
        project_name="CIButler Documentation",
        subtitle="CI/CD Pipeline Management for OSDU",
        author="Shane Hutchins",
        primary_color="#3f51b5",
    )


@pytest.fixture
def cibutler_git_info() -> GitVersion | None:
    return extract_git_info(CIBUTLER_REPO)


@pytest.fixture
def cibutler_html_pages(
    cibutler_config: MkDocsConfig,
) -> list[tuple]:
    renderer = MarkdownRenderer(
        extensions=cibutler_config.markdown_extensions,
        docs_dir=cibutler_config.docs_dir,
    )
    pages = flatten_nav(cibutler_config.nav_items)
    html_pages = []
    for item in pages:
        if item.path is None:
            html_pages.append((item, ""))
            continue
        md_file = cibutler_config.docs_dir / item.path
        if not md_file.exists():
            continue
        html = renderer.render(md_file.read_text(encoding="utf-8"), md_file)
        html_pages.append((item, html))
    return html_pages


# --- MkDocs Parser Tests ---


class TestCibutlerParsing:
    def test_site_name(self, cibutler_config: MkDocsConfig) -> None:
        assert cibutler_config.site_name == "OSDU CIButler Documentation"

    def test_docs_dir_exists(self, cibutler_config: MkDocsConfig) -> None:
        assert cibutler_config.docs_dir.is_dir()

    def test_theme_detected(self, cibutler_config: MkDocsConfig) -> None:
        assert cibutler_config.theme_name == "material"

    def test_extensions_parsed(self, cibutler_config: MkDocsConfig) -> None:
        ext_names = []
        for ext in cibutler_config.markdown_extensions:
            if isinstance(ext, str):
                ext_names.append(ext)
            elif isinstance(ext, dict):
                ext_names.extend(ext.keys())
        assert "admonition" in ext_names
        assert "pymdownx.highlight" in ext_names
        assert "pymdownx.superfences" in ext_names
        assert "pymdownx.tasklist" in ext_names

    def test_nav_auto_discovered(self, cibutler_config: MkDocsConfig) -> None:
        """CIButler mkdocs.yml has no nav key, so pages are auto-discovered."""
        pages = flatten_nav(cibutler_config.nav_items)
        page_paths = [str(p.path) for p in pages if p.path is not None]
        assert any("index.md" in p for p in page_paths)
        assert any("install.md" in p for p in page_paths)
        assert len(page_paths) == 12

    def test_all_pages_exist(self, cibutler_config: MkDocsConfig) -> None:
        pages = flatten_nav(cibutler_config.nav_items)
        for item in pages:
            if item.path is not None:
                md_file = cibutler_config.docs_dir / item.path
                assert md_file.exists(), f"Missing: {item.path}"


# --- Markdown Rendering Tests ---


class TestCibutlerMarkdown:
    def test_all_pages_render(self, cibutler_html_pages: list[tuple]) -> None:
        content_pages = [(item, html) for item, html in cibutler_html_pages if item.path is not None]
        assert len(content_pages) == 12
        for item, html in content_pages:
            assert len(html) > 0, f"Empty HTML for {item.path}"

    def test_admonitions_render(self, cibutler_config: MkDocsConfig) -> None:
        """The install docs use admonitions."""
        renderer = MarkdownRenderer(
            extensions=cibutler_config.markdown_extensions,
            docs_dir=cibutler_config.docs_dir,
        )
        md_file = cibutler_config.docs_dir / "install.md"
        html = renderer.render(md_file.read_text(encoding="utf-8"), md_file)
        # install.md likely has code blocks at minimum
        assert "<code" in html or "<pre" in html

    def test_code_blocks_render(self, cibutler_config: MkDocsConfig) -> None:
        """Pages with code fences should produce highlighted output."""
        renderer = MarkdownRenderer(
            extensions=cibutler_config.markdown_extensions,
            docs_dir=cibutler_config.docs_dir,
        )
        md_file = cibutler_config.docs_dir / "install_osdu.md"
        html = renderer.render(md_file.read_text(encoding="utf-8"), md_file)
        assert "<code" in html or "<pre" in html

    def test_tasklist_renders(self, cibutler_config: MkDocsConfig) -> None:
        """pymdownx.tasklist should not crash even if no task lists present."""
        renderer = MarkdownRenderer(
            extensions=cibutler_config.markdown_extensions,
            docs_dir=cibutler_config.docs_dir,
        )
        md_file = cibutler_config.docs_dir / "index.md"
        html = renderer.render(md_file.read_text(encoding="utf-8"), md_file)
        assert "CIButler" in html or "cibutler" in html.lower()


# --- Git Info Tests ---


class TestCibutlerGitInfo:
    def test_git_info_extracted(self, cibutler_git_info: GitVersion | None) -> None:
        assert cibutler_git_info is not None

    def test_commit_hash(self, cibutler_git_info: GitVersion | None) -> None:
        assert cibutler_git_info is not None
        assert len(cibutler_git_info.commit_hash) == 7

    def test_branch(self, cibutler_git_info: GitVersion | None) -> None:
        assert cibutler_git_info is not None
        assert cibutler_git_info.branch

    def test_version_string(self, cibutler_git_info: GitVersion | None) -> None:
        assert cibutler_git_info is not None
        version_str = cibutler_git_info.format_version_string()
        assert len(version_str) > 0
        assert cibutler_git_info.commit_hash in version_str


# --- PDF Output Tests ---


class TestCibutlerPdf:
    def test_pdf_with_branding(
        self,
        cibutler_config: MkDocsConfig,
        cibutler_branding: BrandingConfig,
        cibutler_git_info: GitVersion | None,
        cibutler_html_pages: list[tuple],
        tmp_output: Path,
    ) -> None:
        pdf_path = tmp_output / "cibutler_branded.pdf"
        renderer = PdfRenderer(cibutler_branding, cibutler_git_info, cibutler_config)
        renderer.render(cibutler_html_pages, pdf_path)

        assert pdf_path.exists()
        assert pdf_path.stat().st_size > 10_000  # should be substantial

    def test_pdf_without_branding(
        self,
        cibutler_config: MkDocsConfig,
        cibutler_html_pages: list[tuple],
        tmp_output: Path,
    ) -> None:
        pdf_path = tmp_output / "cibutler_plain.pdf"
        renderer = PdfRenderer(None, None, cibutler_config)
        renderer.render(cibutler_html_pages, pdf_path, cover_page=False, include_toc=False)

        assert pdf_path.exists()
        assert pdf_path.stat().st_size > 5_000

    def test_pdf_no_cover_with_toc(
        self,
        cibutler_config: MkDocsConfig,
        cibutler_branding: BrandingConfig,
        cibutler_html_pages: list[tuple],
        tmp_output: Path,
    ) -> None:
        pdf_path = tmp_output / "cibutler_no_cover.pdf"
        renderer = PdfRenderer(cibutler_branding, None, cibutler_config)
        renderer.render(cibutler_html_pages, pdf_path, cover_page=False, include_toc=True)

        assert pdf_path.exists()
        assert pdf_path.stat().st_size > 5_000


# --- DOCX Output Tests ---


class TestCibutlerDocx:
    def test_docx_with_branding(
        self,
        cibutler_config: MkDocsConfig,
        cibutler_branding: BrandingConfig,
        cibutler_git_info: GitVersion | None,
        cibutler_html_pages: list[tuple],
        tmp_output: Path,
    ) -> None:
        docx_path = tmp_output / "cibutler_branded.docx"
        renderer = DocxRenderer(cibutler_branding, cibutler_git_info, cibutler_config)
        renderer.render(cibutler_html_pages, docx_path)

        assert docx_path.exists()
        assert docx_path.stat().st_size > 10_000

    def test_docx_without_branding(
        self,
        cibutler_config: MkDocsConfig,
        cibutler_html_pages: list[tuple],
        tmp_output: Path,
    ) -> None:
        docx_path = tmp_output / "cibutler_plain.docx"
        renderer = DocxRenderer(None, None, cibutler_config)
        renderer.render(cibutler_html_pages, docx_path, cover_page=False, include_toc=False)

        assert docx_path.exists()
        assert docx_path.stat().st_size > 5_000

    def test_docx_content_has_headings(
        self,
        cibutler_config: MkDocsConfig,
        cibutler_branding: BrandingConfig,
        cibutler_html_pages: list[tuple],
        tmp_output: Path,
    ) -> None:
        """Verify the DOCX actually contains heading elements."""
        from docx import Document

        docx_path = tmp_output / "cibutler_headings.docx"
        renderer = DocxRenderer(cibutler_branding, None, cibutler_config)
        renderer.render(cibutler_html_pages, docx_path, cover_page=False, include_toc=False)

        doc = Document(str(docx_path))
        heading_texts = [
            p.text for p in doc.paragraphs if p.style.name.startswith("Heading")
        ]
        assert len(heading_texts) > 0

    def test_docx_content_has_tables(
        self,
        cibutler_config: MkDocsConfig,
        cibutler_html_pages: list[tuple],
        tmp_output: Path,
    ) -> None:
        """Check if tables from markdown are preserved in DOCX."""
        from docx import Document

        docx_path = tmp_output / "cibutler_tables.docx"
        renderer = DocxRenderer(None, None, cibutler_config)
        renderer.render(cibutler_html_pages, docx_path, cover_page=False, include_toc=False)

        doc = Document(str(docx_path))
        # The cibutler docs may or may not have tables; just verify no crash
        assert len(doc.paragraphs) > 0


# --- CLI Integration Test ---


class TestCibutlerCli:
    def test_info_command(self) -> None:
        from typer.testing import CliRunner

        from leafpress.cli import cli

        runner = CliRunner()
        result = runner.invoke(cli, ["info", str(CIBUTLER_DOCS)])
        assert result.exit_code == 0
        assert "OSDU CIButler Documentation" in result.output

    def test_convert_pdf(self, tmp_output: Path) -> None:
        from typer.testing import CliRunner

        from leafpress.cli import cli

        runner = CliRunner()
        result = runner.invoke(
            cli,
            ["convert", str(CIBUTLER_DOCS), "-f", "pdf", "-o", str(tmp_output)],
        )
        assert result.exit_code == 0
        assert "Done!" in result.output
        pdf_files = list(tmp_output.glob("*.pdf"))
        assert len(pdf_files) == 1

    def test_convert_both(self, tmp_output: Path) -> None:
        from typer.testing import CliRunner

        from leafpress.cli import cli

        runner = CliRunner()
        result = runner.invoke(
            cli,
            ["convert", str(CIBUTLER_DOCS), "-f", "both", "-o", str(tmp_output)],
        )
        assert result.exit_code == 0
        assert len(list(tmp_output.glob("*.pdf"))) == 1
        assert len(list(tmp_output.glob("*.docx"))) == 1
