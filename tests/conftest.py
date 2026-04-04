"""Shared test fixtures."""

from __future__ import annotations

import struct
import zlib
from pathlib import Path

import pytest
from docx import Document as DocxDocument
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

FIXTURES_DIR = Path(__file__).parent / "fixtures"
SAMPLE_MKDOCS_DIR = FIXTURES_DIR / "sample_mkdocs_project"
SAMPLE_CONFIG_PATH = FIXTURES_DIR / "sample_config.yml"


def _make_png() -> bytes:
    """Create a minimal valid 1x1 red PNG."""

    def _chunk(chunk_type: bytes, data: bytes) -> bytes:
        c = chunk_type + data
        return struct.pack(">I", len(data)) + c + struct.pack(">I", zlib.crc32(c) & 0xFFFFFFFF)

    return (
        b"\x89PNG\r\n\x1a\n"
        + _chunk(b"IHDR", struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0))
        + _chunk(b"IDAT", zlib.compress(b"\x00\xff\x00\x00"))
        + _chunk(b"IEND", b"")
    )


def _add_hyperlink(paragraph, url: str, text: str) -> None:
    """Add a hyperlink to a python-docx paragraph via the oxml API."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run_el = OxmlElement("w:r")
    text_el = OxmlElement("w:t")
    text_el.text = text
    run_el.append(text_el)
    hyperlink.append(run_el)
    paragraph._element.append(hyperlink)


@pytest.fixture
def sample_mkdocs_dir() -> Path:
    return SAMPLE_MKDOCS_DIR


@pytest.fixture
def sample_mkdocs_config() -> Path:
    return SAMPLE_MKDOCS_DIR / "mkdocs.yml"


@pytest.fixture
def sample_branding_config() -> Path:
    return SAMPLE_CONFIG_PATH


@pytest.fixture
def tmp_output(tmp_path: Path) -> Path:
    output = tmp_path / "output"
    output.mkdir()
    return output


@pytest.fixture
def sample_docx(tmp_path: Path) -> Path:
    """Create a minimal DOCX with various formatting for import tests."""
    doc = DocxDocument()
    doc.add_heading("Test Document", level=1)
    p = doc.add_paragraph()
    run_bold = p.add_run("bold text")
    run_bold.bold = True
    p.add_run(" and ")
    run_italic = p.add_run("italic text")
    run_italic.italic = True
    p.add_run(".")

    doc.add_heading("Section Two", level=2)
    doc.add_paragraph("First item", style="List Bullet")
    doc.add_paragraph("Second item", style="List Bullet")
    doc.add_paragraph("Numbered one", style="List Number")
    doc.add_paragraph("Numbered two", style="List Number")

    doc.add_heading("Tables", level=2)
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Header 1"
    table.cell(0, 1).text = "Header 2"
    table.cell(1, 0).text = "Cell A"
    table.cell(1, 1).text = "Cell B"

    path = tmp_path / "test.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def comprehensive_docx(tmp_path: Path) -> Path:
    """Create a feature-rich DOCX with headings, formatting, links, lists, tables, images."""
    doc = DocxDocument()

    # Title and body text
    doc.add_heading("Quarterly Report", level=1)
    doc.add_paragraph(
        "This document summarizes the key results from Q4 2024. "
        "All metrics are compared against Q3 baseline figures."
    )

    # H2 with formatting
    doc.add_heading("Executive Summary", level=2)
    p = doc.add_paragraph()
    run_b = p.add_run("Revenue grew 15%")
    run_b.bold = True
    p.add_run(" while ")
    run_i = p.add_run("operating costs")
    run_i.italic = True
    p.add_run(" remained flat year-over-year.")

    # H3 heading
    doc.add_heading("Regional Breakdown", level=3)
    doc.add_paragraph("North America led growth at 22%, followed by EMEA at 12%.")

    # Hyperlink
    p = doc.add_paragraph("For details, visit ")
    _add_hyperlink(p, "https://example.com/report", "the full report")
    p.add_run(".")

    # Bullet list
    doc.add_heading("Key Highlights", level=2)
    doc.add_paragraph("Customer retention improved to 94%", style="List Bullet")
    doc.add_paragraph("New product launch exceeded targets", style="List Bullet")
    doc.add_paragraph("Headcount grew by 50 engineers", style="List Bullet")

    # Numbered list
    doc.add_heading("Action Items", level=2)
    doc.add_paragraph("Finalize budget proposal", style="List Number")
    doc.add_paragraph("Schedule board review", style="List Number")
    doc.add_paragraph("Update investor deck", style="List Number")

    # Table with 3 columns
    doc.add_heading("Financial Summary", level=2)
    table = doc.add_table(rows=4, cols=3)
    headers = ["Metric", "Q3 2024", "Q4 2024"]
    for i, h in enumerate(headers):
        table.cell(0, i).text = h
    data = [
        ["Revenue", "$2.1M", "$2.4M"],
        ["Expenses", "$1.8M", "$1.8M"],
        ["Net Income", "$300K", "$600K"],
    ]
    for r, row in enumerate(data, start=1):
        for c, val in enumerate(row):
            table.cell(r, c).text = val

    # Image
    doc.add_heading("Chart", level=2)
    img_path = tmp_path / "chart.png"
    img_path.write_bytes(_make_png())
    doc.add_picture(str(img_path))

    path = tmp_path / "quarterly_report.docx"
    doc.save(str(path))
    return path
