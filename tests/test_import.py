"""Tests for DOCX to Markdown import."""

from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest
from docx import Document as DocxDocument
from typer.testing import CliRunner

from leafpress.cli import _resolve_import_source, cli
from leafpress.exceptions import DocxImportError, LeafpressError
from leafpress.importer.base import ImportResult, postprocess_markdown
from leafpress.importer.converter import import_docx

runner = CliRunner()


def test_import_basic(sample_docx: Path, tmp_output: Path) -> None:
    """End-to-end: DOCX in, .md out, file exists with content."""
    result = import_docx(sample_docx, output_path=tmp_output)
    assert result.markdown_path.exists()
    assert result.markdown_path.suffix == ".md"
    content = result.markdown_path.read_text()
    assert len(content) > 0
    assert "Test Document" in content


def test_import_headings(sample_docx: Path, tmp_output: Path) -> None:
    """Headings converted to ATX-style markdown."""
    result = import_docx(sample_docx, output_path=tmp_output)
    content = result.markdown_path.read_text()
    assert "# Test Document" in content
    assert "## Section Two" in content
    assert "## Tables" in content


def test_import_formatting(sample_docx: Path, tmp_output: Path) -> None:
    """Bold and italic formatting preserved."""
    result = import_docx(sample_docx, output_path=tmp_output)
    content = result.markdown_path.read_text()
    assert "**bold text**" in content
    assert "*italic text*" in content


def test_import_lists(sample_docx: Path, tmp_output: Path) -> None:
    """Ordered and unordered lists present in output."""
    result = import_docx(sample_docx, output_path=tmp_output)
    content = result.markdown_path.read_text()
    assert "First item" in content
    assert "Second item" in content
    assert "Numbered one" in content
    assert "Numbered two" in content


def test_import_tables(sample_docx: Path, tmp_output: Path) -> None:
    """Tables converted to pipe-style markdown."""
    result = import_docx(sample_docx, output_path=tmp_output)
    content = result.markdown_path.read_text()
    assert "Header 1" in content
    assert "Header 2" in content
    assert "Cell A" in content
    assert "Cell B" in content
    # Should have pipe table separator
    assert "---" in content


def test_import_images_extracted(tmp_path: Path, tmp_output: Path) -> None:
    """Images extracted to assets/ and referenced in markdown."""
    # Create a DOCX with an image
    doc = DocxDocument()
    doc.add_heading("Image Test", level=1)
    # Create a small PNG (1x1 pixel)
    import struct
    import zlib

    def _make_png() -> bytes:
        """Create minimal 1x1 red PNG."""

        def _chunk(chunk_type: bytes, data: bytes) -> bytes:
            c = chunk_type + data
            return struct.pack(">I", len(data)) + c + struct.pack(">I", zlib.crc32(c) & 0xFFFFFFFF)

        sig = b"\x89PNG\r\n\x1a\n"
        ihdr = _chunk(b"IHDR", struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0))
        raw = zlib.compress(b"\x00\xff\x00\x00")
        idat = _chunk(b"IDAT", raw)
        iend = _chunk(b"IEND", b"")
        return sig + ihdr + idat + iend

    img_path = tmp_path / "test.png"
    img_path.write_bytes(_make_png())
    doc.add_picture(str(img_path))
    docx_path = tmp_path / "with_image.docx"
    doc.save(str(docx_path))

    result = import_docx(docx_path, output_path=tmp_output)
    assert len(result.images) >= 1
    assert result.images[0].exists()
    content = result.markdown_path.read_text()
    assert "assets/" in content


def test_import_no_images_flag(tmp_path: Path, tmp_output: Path) -> None:
    """--no-extract-images skips image extraction."""
    doc = DocxDocument()
    doc.add_heading("No Images", level=1)
    docx_path = tmp_path / "no_images.docx"
    doc.save(str(docx_path))

    result = import_docx(docx_path, output_path=tmp_output, extract_images=False)
    assert result.images == []
    assets_dir = tmp_output / "assets"
    assert not assets_dir.exists()


def test_import_output_dir(sample_docx: Path, tmp_output: Path) -> None:
    """-o dir/ creates dir/stem.md."""
    result = import_docx(sample_docx, output_path=tmp_output)
    assert result.markdown_path == tmp_output / "test.md"
    assert result.markdown_path.exists()


def test_import_output_file(sample_docx: Path, tmp_output: Path) -> None:
    """-o file.md writes to exact path."""
    out_file = tmp_output / "custom.md"
    result = import_docx(sample_docx, output_path=out_file)
    assert result.markdown_path == out_file
    assert out_file.exists()


def test_import_default_output(sample_docx: Path) -> None:
    """No -o writes sibling .md next to the docx."""
    result = import_docx(sample_docx, output_path=None)
    expected = sample_docx.with_suffix(".md")
    assert result.markdown_path == expected
    assert expected.exists()


def test_import_nonexistent_file(tmp_path: Path) -> None:
    """Raises DocxImportError for missing file."""
    with pytest.raises(DocxImportError, match="File not found"):
        import_docx(tmp_path / "nonexistent.docx")


def test_import_non_docx_file(tmp_path: Path) -> None:
    """Raises DocxImportError for non-.docx file."""
    txt_file = tmp_path / "readme.txt"
    txt_file.write_text("hello")
    with pytest.raises(DocxImportError, match=r"Not a \.docx file"):
        import_docx(txt_file)


def test_import_cli(sample_docx: Path, tmp_output: Path) -> None:
    """CLI end-to-end test."""
    result = runner.invoke(cli, ["import", str(sample_docx), "-o", str(tmp_output)])
    assert result.exit_code == 0
    assert "Done!" in result.output
    assert (tmp_output / "test.md").exists()


def test_import_cli_nonexistent(tmp_path: Path) -> None:
    """CLI exits 1 for missing file."""
    result = runner.invoke(cli, ["import", str(tmp_path / "missing.docx")])
    assert result.exit_code == 1
    assert "Error" in result.output


# ---------------------------------------------------------------------------
# URL import source resolution tests
# ---------------------------------------------------------------------------


def test_resolve_local_path(sample_docx: Path) -> None:
    """Local paths pass through unchanged."""
    with _resolve_import_source(str(sample_docx)) as resolved:
        assert resolved == sample_docx


def test_resolve_url_downloads_file() -> None:
    """URLs are downloaded to a temp file with the correct extension."""
    mock_resp = MagicMock()
    mock_resp.headers = {"content-type": "application/octet-stream"}
    mock_resp.iter_content.return_value = [b"fake docx content"]
    mock_resp.raise_for_status = MagicMock()

    with (
        patch("leafpress.cli.requests.get", return_value=mock_resp),
        _resolve_import_source("https://example.com/report.docx") as resolved,
    ):
        assert resolved.suffix == ".docx"
        assert resolved.stem == "report"
        assert resolved.exists()
        assert resolved.read_bytes() == b"fake docx content"


def test_resolve_url_infers_tex_from_extension() -> None:
    """URL with .tex extension is recognized."""
    mock_resp = MagicMock()
    mock_resp.headers = {"content-type": "text/plain"}
    mock_resp.iter_content.return_value = [b"\\documentclass{article}"]
    mock_resp.raise_for_status = MagicMock()

    with (
        patch("leafpress.cli.requests.get", return_value=mock_resp),
        _resolve_import_source("https://example.com/paper.tex") as resolved,
    ):
        assert resolved.suffix == ".tex"


def test_resolve_url_falls_back_to_content_type() -> None:
    """URLs without extension use Content-Type to determine format."""
    mock_resp = MagicMock()
    mock_resp.headers = {
        "content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    }
    mock_resp.iter_content.return_value = [b"fake"]
    mock_resp.raise_for_status = MagicMock()

    with (
        patch("leafpress.cli.requests.get", return_value=mock_resp),
        _resolve_import_source("https://example.com/download") as resolved,
    ):
        assert resolved.suffix == ".docx"


def test_resolve_url_unknown_type_raises() -> None:
    """URLs with no extension and unrecognized Content-Type raise an error."""
    mock_resp = MagicMock()
    mock_resp.headers = {"content-type": "text/html"}
    mock_resp.iter_content.return_value = [b"<html>"]
    mock_resp.raise_for_status = MagicMock()

    with (
        patch("leafpress.cli.requests.get", return_value=mock_resp),
        pytest.raises(LeafpressError, match="Cannot determine file type"),
        _resolve_import_source("https://example.com/page"),
    ):
        pass


def test_resolve_url_http_error_raises() -> None:
    """HTTP errors are wrapped in LeafpressError."""
    import requests

    mock_resp = MagicMock()
    mock_resp.raise_for_status.side_effect = requests.HTTPError("404 Not Found")

    with (
        patch("leafpress.cli.requests.get", return_value=mock_resp),
        pytest.raises(LeafpressError, match="Failed to download"),
        _resolve_import_source("https://example.com/missing.docx"),
    ):
        pass


def test_resolve_url_temp_dir_cleaned_up() -> None:
    """Temp directory is cleaned up after the context manager exits."""
    mock_resp = MagicMock()
    mock_resp.headers = {"content-type": "application/octet-stream"}
    mock_resp.iter_content.return_value = [b"data"]
    mock_resp.raise_for_status = MagicMock()

    with patch("leafpress.cli.requests.get", return_value=mock_resp):
        with _resolve_import_source("https://example.com/file.docx") as resolved:
            temp_dir = resolved.parent
            assert temp_dir.exists()
        # After context exit, temp dir should be gone
        assert not temp_dir.exists()


def test_import_cli_with_url(sample_docx: Path, tmp_output: Path) -> None:
    """CLI accepts a URL and imports the downloaded file."""
    # Serve the real sample_docx content via a mocked URL
    docx_bytes = sample_docx.read_bytes()
    mock_resp = MagicMock()
    mock_resp.headers = {
        "content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    }
    mock_resp.iter_content.return_value = [docx_bytes]
    mock_resp.raise_for_status = MagicMock()

    with patch("leafpress.cli.requests.get", return_value=mock_resp):
        result = runner.invoke(
            cli,
            ["import", "https://example.com/test.docx", "-o", str(tmp_output)],
        )
    assert result.exit_code == 0
    assert "Done!" in result.output
    assert (tmp_output / "test.md").exists()


def test_postprocess_markdown() -> None:
    """Blank lines collapsed, trailing whitespace stripped."""
    raw = "# Title\n\n\n\nParagraph   \n\n\n\n\nEnd  "
    cleaned = postprocess_markdown(raw)
    assert "\n\n\n" not in cleaned
    assert "   \n" not in cleaned
    assert cleaned.endswith("\n")
    assert not cleaned.endswith("\n\n")


def test_import_result_dataclass() -> None:
    """ImportResult holds expected fields."""
    r = ImportResult(markdown_path=Path("test.md"))
    assert r.markdown_path == Path("test.md")
    assert r.images == []
    assert r.warnings == []


# --- Direct markdown converter unit tests ---

from leafpress.importer.markdown_converter import LeafpressMarkdownConverter  # noqa: E402


def _convert_html(html: str) -> str:
    """Helper to convert HTML snippet through the converter."""
    return LeafpressMarkdownConverter(strip=["html", "body"]).convert(html).strip()


def test_converter_fenced_code_block() -> None:
    """<pre><code> converts to fenced code block."""
    html = "<pre><code>print('hello')</code></pre>"
    md = _convert_html(html)
    assert "```" in md
    assert "print('hello')" in md


def test_converter_fenced_code_with_language() -> None:
    """<pre><code class='language-python'> gets language annotation."""
    html = '<pre><code class="language-python">x = 1</code></pre>'
    md = _convert_html(html)
    assert "```python" in md
    assert "x = 1" in md


def test_converter_fenced_code_no_language() -> None:
    """<pre><code> without language class produces bare fence."""
    html = "<pre><code>some code</code></pre>"
    md = _convert_html(html)
    assert "```\n" in md
    assert "some code" in md


def test_converter_fenced_code_non_language_class() -> None:
    """<pre><code class='highlight'> (not language-*) produces bare fence."""
    html = '<pre><code class="highlight">foo</code></pre>'
    md = _convert_html(html)
    assert "```\n" in md or md.startswith("```\n") or "```\nfoo" in md


def test_converter_pipe_table() -> None:
    """<table> converts to pipe-style markdown table."""
    html = "<table><tr><th>A</th><th>B</th></tr><tr><td>1</td><td>2</td></tr></table>"
    md = _convert_html(html)
    assert "| A" in md
    assert "| B" in md
    assert "---" in md
    assert "| 1" in md
    assert "| 2" in md


def test_converter_table_uneven_rows() -> None:
    """Table with uneven row lengths pads shorter rows."""
    html = "<table><tr><th>A</th><th>B</th><th>C</th></tr><tr><td>1</td><td>2</td></tr></table>"
    md = _convert_html(html)
    # Should still produce a valid table with 3 columns
    lines = [line for line in md.split("\n") if line.startswith("|")]
    assert len(lines) == 3  # header + separator + data row
    # Each line should have 3 pipe-separated cells
    for line in lines:
        assert line.count("|") >= 4  # | cell | cell | cell |


def test_converter_table_empty() -> None:
    """<table> with no rows falls through gracefully."""
    html = "<table></table>"
    md = _convert_html(html)
    # Should not crash; result is whatever fallback markdownify gives
    assert isinstance(md, str)


# ---------------------------------------------------------------------------
# Comprehensive DOCX fixture tests — realistic multi-feature document
# ---------------------------------------------------------------------------


class TestComprehensiveDocx:
    """Tests using the comprehensive_docx fixture (realistic business document)."""

    @pytest.fixture(autouse=True)
    def _convert(self, comprehensive_docx: Path, tmp_output: Path) -> None:
        self.result = import_docx(comprehensive_docx, output_path=tmp_output)
        self.content = self.result.markdown_path.read_text()

    def test_converts_successfully(self) -> None:
        """Full document converts without errors."""
        assert self.result.markdown_path.exists()
        assert len(self.content) > 200

    def test_title_heading(self) -> None:
        """H1 title present."""
        assert "# Quarterly Report" in self.content

    def test_h2_headings(self) -> None:
        """H2 headings present."""
        assert "## Executive Summary" in self.content
        assert "## Key Highlights" in self.content
        assert "## Financial Summary" in self.content

    def test_h3_heading(self) -> None:
        """H3 heading present."""
        assert "### Regional Breakdown" in self.content

    def test_bold_formatting(self) -> None:
        """Bold text preserved."""
        assert "**Revenue grew 15%**" in self.content

    def test_italic_formatting(self) -> None:
        """Italic text preserved."""
        assert "*operating costs*" in self.content

    def test_hyperlink(self) -> None:
        """Hyperlink converted to markdown link syntax."""
        assert "[the full report](https://example.com/report)" in self.content

    def test_bullet_list(self) -> None:
        """Bullet list items present."""
        assert "Customer retention" in self.content
        assert "New product launch" in self.content
        assert "Headcount grew" in self.content

    def test_numbered_list(self) -> None:
        """Numbered list items present."""
        assert "Finalize budget" in self.content
        assert "Schedule board" in self.content
        assert "Update investor" in self.content

    def test_table(self) -> None:
        """Table converted to pipe format with all data."""
        assert "Metric" in self.content
        assert "Q3 2024" in self.content
        assert "Revenue" in self.content
        assert "$2.4M" in self.content
        assert "---" in self.content

    def test_image_extracted(self) -> None:
        """Embedded image extracted to assets/."""
        assert len(self.result.images) >= 1
        assert self.result.images[0].exists()
        assert "assets/" in self.content

    def test_body_text(self) -> None:
        """Body paragraphs preserved."""
        assert "Q4 2024" in self.content
        assert "North America led growth" in self.content
