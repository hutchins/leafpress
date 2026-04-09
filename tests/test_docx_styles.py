"""Tests for DOCX branding styles."""

from typing import Any

from docx import Document
from docx.shared import Pt, RGBColor

from leafpress.config import BrandingConfig
from leafpress.docx.styles import _parse_hex_color, apply_branding_styles


def _make_branding(**kwargs: Any) -> BrandingConfig:
    defaults: dict[str, Any] = {"company_name": "TestCo", "project_name": "TestProject"}
    defaults.update(kwargs)
    return BrandingConfig(**defaults)


def test_apply_branding_styles() -> None:
    doc = Document()
    branding = _make_branding(primary_color="#2e7d32")
    apply_branding_styles(doc, branding)

    normal = doc.styles["Normal"]
    assert normal.font.name == "Calibri"
    assert normal.font.size == Pt(11)

    h1 = doc.styles["Heading 1"]
    assert h1.font.size == Pt(24)
    assert h1.font.color.rgb == RGBColor(0x2E, 0x7D, 0x32)


def test_apply_styles_without_branding() -> None:
    doc = Document()
    apply_branding_styles(doc, None)

    # Uses default color #1a73e8
    h1 = doc.styles["Heading 1"]
    assert h1.font.color.rgb == RGBColor(0x1A, 0x73, 0xE8)


def test_heading_levels() -> None:
    doc = Document()
    apply_branding_styles(doc, _make_branding())

    assert doc.styles["Heading 1"].font.size == Pt(24)
    assert doc.styles["Heading 2"].font.size == Pt(18)
    assert doc.styles["Heading 3"].font.size == Pt(14)
    assert doc.styles["Heading 4"].font.size == Pt(12)


def test_parse_hex_color() -> None:
    assert _parse_hex_color("#1a73e8") == RGBColor(0x1A, 0x73, 0xE8)
    assert _parse_hex_color("#FF0000") == RGBColor(0xFF, 0x00, 0x00)
    assert _parse_hex_color("#000000") == RGBColor(0x00, 0x00, 0x00)


def test_parse_hex_color_invalid() -> None:
    import pytest

    with pytest.raises(ValueError, match="Invalid hex color"):
        _parse_hex_color("#fff")
    with pytest.raises(ValueError, match="Invalid hex color"):
        _parse_hex_color("invalid")
