"""Tests for DOCX HTML-to-document converter."""

from pathlib import Path

from docx import Document

from leafpress.docx.html_converter import HtmlToDocxConverter


def _make_converter(tmp_path: Path) -> tuple[Document, HtmlToDocxConverter]:
    doc = Document()
    converter = HtmlToDocxConverter(doc, tmp_path)
    return doc, converter


def test_convert_empty_html(tmp_path: Path) -> None:
    doc, converter = _make_converter(tmp_path)
    converter.convert("")
    converter.convert("   ")
    # No paragraphs added beyond the default
    assert len(doc.paragraphs) == 0


def test_convert_heading(tmp_path: Path) -> None:
    doc, converter = _make_converter(tmp_path)
    converter.convert("<h1>Title</h1>")
    converter.convert("<h2>Subtitle</h2>")
    converter.convert("<h3>Section</h3>")
    assert any(p.text == "Title" for p in doc.paragraphs)
    assert any(p.text == "Subtitle" for p in doc.paragraphs)
    assert any(p.text == "Section" for p in doc.paragraphs)


def test_convert_paragraph(tmp_path: Path) -> None:
    doc, converter = _make_converter(tmp_path)
    converter.convert("<p>Hello world</p>")
    assert any("Hello world" in p.text for p in doc.paragraphs)


def test_convert_bold_italic(tmp_path: Path) -> None:
    doc, converter = _make_converter(tmp_path)
    converter.convert("<p><strong>bold</strong> and <em>italic</em></p>")
    para = doc.paragraphs[0]
    runs = para.runs
    bold_runs = [r for r in runs if r.bold]
    italic_runs = [r for r in runs if r.italic]
    assert len(bold_runs) >= 1
    assert len(italic_runs) >= 1


def test_convert_inline_code(tmp_path: Path) -> None:
    doc, converter = _make_converter(tmp_path)
    converter.convert("<p>Use <code>print()</code> here</p>")
    para = doc.paragraphs[0]
    code_runs = [r for r in para.runs if r.font.name == "Courier New"]
    assert len(code_runs) >= 1


def test_convert_link(tmp_path: Path) -> None:
    doc, converter = _make_converter(tmp_path)
    converter.convert('<p><a href="https://example.com">click</a></p>')
    para = doc.paragraphs[0]
    link_runs = [r for r in para.runs if r.underline]
    assert len(link_runs) >= 1


def test_convert_unordered_list(tmp_path: Path) -> None:
    doc, converter = _make_converter(tmp_path)
    converter.convert("<ul><li>One</li><li>Two</li></ul>")
    texts = [p.text for p in doc.paragraphs]
    assert "One" in texts
    assert "Two" in texts


def test_convert_ordered_list(tmp_path: Path) -> None:
    doc, converter = _make_converter(tmp_path)
    converter.convert("<ol><li>First</li><li>Second</li></ol>")
    texts = [p.text for p in doc.paragraphs]
    assert "First" in texts
    assert "Second" in texts


def test_convert_table(tmp_path: Path) -> None:
    doc, converter = _make_converter(tmp_path)
    converter.convert("<table><tr><th>Header</th></tr><tr><td>Cell</td></tr></table>")
    assert len(doc.tables) == 1
    table = doc.tables[0]
    assert table.cell(0, 0).text == "Header"
    assert table.cell(1, 0).text == "Cell"


def test_convert_code_block(tmp_path: Path) -> None:
    doc, converter = _make_converter(tmp_path)
    converter.convert("<pre><code>x = 1</code></pre>")
    code_paras = [p for p in doc.paragraphs if p.runs and p.runs[0].font.name == "Courier New"]
    assert len(code_paras) >= 1


def test_convert_blockquote(tmp_path: Path) -> None:
    doc, converter = _make_converter(tmp_path)
    converter.convert("<blockquote><p>A quote</p></blockquote>")
    assert any("A quote" in p.text for p in doc.paragraphs)


def test_convert_horizontal_rule(tmp_path: Path) -> None:
    doc, converter = _make_converter(tmp_path)
    converter.convert("<hr>")
    # HR adds a paragraph with a border element
    assert len(doc.paragraphs) >= 1


def test_convert_admonition(tmp_path: Path) -> None:
    doc, converter = _make_converter(tmp_path)
    converter.convert(
        '<div class="admonition note"><p class="admonition-title">Note</p><p>Content here</p></div>'
    )
    texts = [p.text for p in doc.paragraphs]
    assert "Note" in texts
    assert any("Content here" in t for t in texts)


def test_convert_definition_list(tmp_path: Path) -> None:
    doc, converter = _make_converter(tmp_path)
    converter.convert("<dl><dt>Term</dt><dd>Definition</dd></dl>")
    texts = [p.text for p in doc.paragraphs]
    assert "Term" in texts
    assert any("Definition" in t for t in texts)


def test_convert_details(tmp_path: Path) -> None:
    doc, converter = _make_converter(tmp_path)
    converter.convert("<details><summary>More info</summary><p>Details here</p></details>")
    texts = [p.text for p in doc.paragraphs]
    assert "More info" in texts


def test_convert_nested_div(tmp_path: Path) -> None:
    doc, converter = _make_converter(tmp_path)
    converter.convert("<div><p>Inside div</p></div>")
    assert any("Inside div" in p.text for p in doc.paragraphs)


def test_convert_image_not_found(tmp_path: Path) -> None:
    doc, converter = _make_converter(tmp_path)
    converter.convert('<img src="nonexistent.png">')
    assert any("Image not found" in p.text for p in doc.paragraphs)


def test_admonition_colors(tmp_path: Path) -> None:
    _, converter = _make_converter(tmp_path)
    # Test various admonition types return different colors
    note_color = converter._admonition_color(["admonition", "note"])
    warning_color = converter._admonition_color(["admonition", "warning"])
    tip_color = converter._admonition_color(["admonition", "tip"])
    default_color = converter._admonition_color(["admonition", "unknown"])
    assert note_color != warning_color
    assert note_color != tip_color
    assert default_color == note_color  # default is same as note


def test_convert_empty_table(tmp_path: Path) -> None:
    """Table with no rows or no columns is skipped."""
    doc, converter = _make_converter(tmp_path)
    converter.convert("<table></table>")
    assert len(doc.tables) == 0

    doc2, converter2 = _make_converter(tmp_path)
    converter2.convert("<table><tr></tr></table>")
    assert len(doc2.tables) == 0


def test_convert_image_empty_src(tmp_path: Path) -> None:
    """Image with empty src is skipped."""
    doc, converter = _make_converter(tmp_path)
    converter.convert('<img src="">')
    # No "Image not found" paragraph added
    assert len(doc.paragraphs) == 0


def test_convert_image_file_uri(tmp_path: Path) -> None:
    """Image with file:// URI resolves to filesystem path."""
    img_path = tmp_path / "photo.png"
    # Create a minimal valid PNG (1x1 white pixel)
    import struct
    import zlib

    def _make_png() -> bytes:
        sig = b"\x89PNG\r\n\x1a\n"
        ihdr_data = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
        ihdr_crc = struct.pack(">I", zlib.crc32(b"IHDR" + ihdr_data) & 0xFFFFFFFF)
        ihdr = struct.pack(">I", 13) + b"IHDR" + ihdr_data + ihdr_crc
        raw = b"\x00\x00\x00\x00"  # filter=None, R=0, G=0, B=0
        idat_data = zlib.compress(raw)
        idat_crc = struct.pack(">I", zlib.crc32(b"IDAT" + idat_data) & 0xFFFFFFFF)
        idat = struct.pack(">I", len(idat_data)) + b"IDAT" + idat_data + idat_crc
        iend_crc = struct.pack(">I", zlib.crc32(b"IEND") & 0xFFFFFFFF)
        iend = struct.pack(">I", 0) + b"IEND" + iend_crc
        return sig + ihdr + idat + iend

    img_path.write_bytes(_make_png())
    doc, converter = _make_converter(tmp_path)
    converter.convert(f'<img src="file://{img_path}">')
    # Image should be embedded (no "not found" text)
    assert not any("not found" in p.text for p in doc.paragraphs)


def test_convert_image_embed_failure(tmp_path: Path) -> None:
    """Image that exists but fails to embed shows fallback text."""
    bad_img = tmp_path / "bad.png"
    bad_img.write_bytes(b"not a real image")
    doc, converter = _make_converter(tmp_path)
    import warnings

    with warnings.catch_warnings():
        warnings.simplefilter("ignore")
        converter.convert('<img src="bad.png">')
    assert any("[Image:" in p.text for p in doc.paragraphs)


def test_convert_strikethrough(tmp_path: Path) -> None:
    """<del> tag applies strikethrough formatting."""
    doc, converter = _make_converter(tmp_path)
    converter.convert("<p><del>removed</del></p>")
    para = doc.paragraphs[0]
    strike_runs = [r for r in para.runs if r.font.strike]
    assert len(strike_runs) >= 1
    assert any("removed" in r.text for r in strike_runs)


def test_convert_superscript(tmp_path: Path) -> None:
    """<sup> tag applies superscript formatting."""
    doc, converter = _make_converter(tmp_path)
    converter.convert("<p>x<sup>2</sup></p>")
    para = doc.paragraphs[0]
    sup_runs = [r for r in para.runs if r.font.superscript]
    assert len(sup_runs) >= 1
    assert any("2" in r.text for r in sup_runs)


def test_convert_subscript(tmp_path: Path) -> None:
    """<sub> tag applies subscript formatting."""
    doc, converter = _make_converter(tmp_path)
    converter.convert("<p>H<sub>2</sub>O</p>")
    para = doc.paragraphs[0]
    sub_runs = [r for r in para.runs if r.font.subscript]
    assert len(sub_runs) >= 1
    assert any("2" in r.text for r in sub_runs)


def test_convert_mark_highlight(tmp_path: Path) -> None:
    """<mark> tag applies yellow highlight."""
    from docx.enum.text import WD_COLOR_INDEX

    doc, converter = _make_converter(tmp_path)
    converter.convert("<p><mark>highlighted</mark></p>")
    para = doc.paragraphs[0]
    hl_runs = [r for r in para.runs if r.font.highlight_color == WD_COLOR_INDEX.YELLOW]
    assert len(hl_runs) >= 1


def test_convert_br_line_break(tmp_path: Path) -> None:
    """<br> inserts a newline character."""
    doc, converter = _make_converter(tmp_path)
    converter.convert("<p>line1<br>line2</p>")
    para = doc.paragraphs[0]
    assert "\n" in para.text


def test_convert_inline_image(tmp_path: Path) -> None:
    """Inline <img> inside a paragraph delegates to image handler."""
    doc, converter = _make_converter(tmp_path)
    converter.convert("<p>See: <img src='missing.png'></p>")
    all_text = " ".join(p.text for p in doc.paragraphs)
    assert "Image not found" in all_text


def test_convert_emoji_image(tmp_path: Path) -> None:
    """Emoji images (twemoji/gemoji) render as alt text."""
    doc, converter = _make_converter(tmp_path)
    converter.convert('<p>Hello <img class="twemoji" alt="👋" src="wave.png"></p>')
    para = doc.paragraphs[0]
    full_text = "".join(r.text for r in para.runs)
    assert "\U0001f44b" in full_text


def test_convert_definition_list_whitespace(tmp_path: Path) -> None:
    """Definition list with whitespace text nodes between dt/dd is handled."""
    doc, converter = _make_converter(tmp_path)
    converter.convert("<dl>\n  <dt>Key</dt>\n  <dd>Value</dd>\n</dl>")
    texts = [p.text for p in doc.paragraphs]
    assert "Key" in texts
    assert any("Value" in t for t in texts)


def test_convert_no_body_tag(tmp_path: Path) -> None:
    """HTML fragment without explicit body tag still converts."""
    doc, converter = _make_converter(tmp_path)
    # lxml always wraps in html/body, but test a bare fragment to be safe
    converter.convert("<p>bare fragment</p>")
    assert any("bare fragment" in p.text for p in doc.paragraphs)


def test_convert_task_list_checked(tmp_path: Path) -> None:
    """Task list with checked item renders checkbox symbol."""
    doc, converter = _make_converter(tmp_path)
    converter.convert(
        '<ul><li class="task-list-item"><input type="checkbox" checked> Done task</li></ul>'
    )
    full_text = " ".join(p.text for p in doc.paragraphs)
    assert "\u2611" in full_text  # ☑


def test_convert_task_list_unchecked(tmp_path: Path) -> None:
    """Task list with unchecked item renders empty checkbox symbol."""
    doc, converter = _make_converter(tmp_path)
    converter.convert(
        '<ul><li class="task-list-item"><input type="checkbox"> Pending task</li></ul>'
    )
    full_text = " ".join(p.text for p in doc.paragraphs)
    assert "\u2610" in full_text  # ☐


def test_convert_task_list_skips_label_and_input(tmp_path: Path) -> None:
    """Task list items skip label and input elements in inline content."""
    doc, converter = _make_converter(tmp_path)
    converter.convert(
        '<ul><li class="task-list-item"><label><input type="checkbox"> My task</label></li></ul>'
    )
    para = doc.paragraphs[0]
    # The checkbox symbol should appear, and label/input should not duplicate
    assert any("\u2610" in r.text for r in para.runs)


def test_convert_nested_list(tmp_path: Path) -> None:
    """Nested lists are rendered with increased indent level."""
    doc, converter = _make_converter(tmp_path)
    converter.convert("<ul><li>Parent  <ul><li>Child</li></ul></li></ul>")
    texts = [p.text.strip() for p in doc.paragraphs]
    assert "Parent" in texts
    assert "Child" in texts
